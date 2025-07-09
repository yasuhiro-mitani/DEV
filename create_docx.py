from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add a title
title = doc.add_heading('サンプルドキュメント', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a subtitle
doc.add_heading('概要', level=1)

# Add some paragraphs
doc.add_paragraph('このドキュメントはpython-docxライブラリを使用して作成されました。')
doc.add_paragraph('Microsoft Wordで開くことができる.docx形式のファイルです。')

# Add a section with bullet points
doc.add_heading('主な特徴', level=1)
doc.add_paragraph('テキストの追加', style='List Bullet')
doc.add_paragraph('見出しの設定', style='List Bullet')
doc.add_paragraph('段落の書式設定', style='List Bullet')
doc.add_paragraph('リストの作成', style='List Bullet')

# Add another section
doc.add_heading('使用方法', level=1)
doc.add_paragraph('1. このファイルをMicrosoft Wordで開く')
doc.add_paragraph('2. 内容を編集する')
doc.add_paragraph('3. 必要に応じて保存する')

# Add a page break
doc.add_page_break()

# Add content on the second page
doc.add_heading('追加情報', level=1)
doc.add_paragraph('このドキュメントは自動生成されたサンプルです。')
doc.add_paragraph('実際の用途に応じて内容を変更してください。')

# Save the document
doc.save('sample_document.docx')
print("sample_document.docx が正常に作成されました。")