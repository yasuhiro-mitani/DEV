from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Create a new Document
doc = Document()

# Read the markdown file
with open('ai_business_presentation_script.md', 'r', encoding='utf-8') as file:
    content = file.read()

# Split content into lines
lines = content.split('\n')

for line in lines:
    line = line.strip()
    
    if not line:
        continue
    
    # Handle main heading (# title)
    if line.startswith('# '):
        title_text = line[2:]
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Handle section headings (## title)
    elif line.startswith('## '):
        heading_text = line[3:]
        doc.add_heading(heading_text, level=1)
    
    # Handle subsection headings (### title)
    elif line.startswith('### '):
        heading_text = line[4:]
        doc.add_heading(heading_text, level=2)
    
    # Handle bullet points (- item)
    elif line.startswith('- '):
        bullet_text = line[2:]
        doc.add_paragraph(bullet_text, style='List Bullet')
    
    # Handle regular paragraphs
    else:
        doc.add_paragraph(line)

# Save the document
doc.save('ai_business_presentation_script.docx')
print("ai_business_presentation_script.docx が正常に作成されました。")